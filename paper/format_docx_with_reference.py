"""Generate a Word copy of main.tex using a reference DOCX format."""

from __future__ import annotations

import argparse
import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


PAPER_DIR = Path(__file__).resolve().parent
DEFAULT_REFERENCE = Path(os.environ.get('PAPER_REFERENCE_DOCX', PAPER_DIR / 'reference.docx')).expanduser()
DEFAULT_OUTPUT = PAPER_DIR / 'main_123_format.docx'

TITLE = '级联视觉智能：基于多阶段筛选与优先级调度的高效社区安全监控'
AUTHOR = '匿名投稿'
EN_ABSTRACT = (
    'Abstract: Vision-language models (VLMs) are powerful for understanding '
    'surveillance frames, but their computational cost makes per-frame real-time '
    'processing impractical on edge hardware. This paper presents a cascaded '
    'inference pipeline for community safety monitoring and evaluates its VLM '
    'call-volume upper bound in an oracle simulation. The three-stage pipeline '
    'combines motion-based adaptive sampling, COCO-SSD object pre-screening, and '
    'priority-based VLM scheduling. Across four synthetic activity scenarios, '
    'the full cascade reduces VLM calls from 3,000 to 169, a 94.4% reduction. '
    'A quantized Qwen3.5-4B GGUF model is deployed locally with llama.cpp, and '
    'component-level latency is measured on consumer hardware. The oracle '
    'scheduled-event coverage upper bound does not represent end-to-end accuracy '
    'in real surveillance scenes.'
)
EN_KEYWORDS = (
    'Keywords: Vision-Language Model; Edge Deployment; Cascaded Inference; '
    'Community Safety Monitoring; Oracle Simulation'
)

SECTION_TITLES = {'引言', '相关工作', '方法', '实验', '讨论', '结论', '参考文献'}
SUBSECTION_TITLES = {
    '系统架构',
    '级联推理流水线',
    '基于优先级的调度',
    'VLM响应解析',
    '实验设置',
    '未评估内容',
    '主要结果',
    '逐场景分析',
    '消融研究',
}

TABLE_CAPTIONS = [
    '与相关工作的对比。',
    '风险类别与COCO-SSD触发条件的关系。',
    '仿真器关键配置。',
    '注入风险事件的时间窗和触发标签。',
    '各场景、各配置的VLM调用次数。',
    '当前证据支持的主张边界。',
    '主要结果：跨流水线配置的VLM调用缩减和oracle调度覆盖上界，四种活动场景的平均值（每个场景5分钟）。',
    '系统部署参数。',
    '逐场景结果。VLM调用缩减随活动水平变化，但即使在高活动场景中也保持在86%以上。覆盖率为oracle调度上界，并非真实检测率。',
    '累积消融研究：每个配置相对于上一行的额外VLM调用缩减。',
]

FIGURE_CAPTIONS = [
    '级联推理流水线。阶段1：自适应帧捕获，过滤静态帧。阶段2：COCO-SSD目标检测预筛选。阶段3：基于优先级的VLM分析——人体检测触发即时调度，其他目标采用机会性调度。',
    '基于优先级的VLM调度。人体检测触发即时调度，中止任何过时的进行中请求。其他目标仅在无进行中请求时采用机会性调度。',
    '跨流水线配置的VLM调用量。完整级联（红色）相比朴素基线（灰色）减少了94.4%的调用。百分比标签显示相对于朴素基线的缩减率。',
    '左图：朴素基线与完整级联在不同活动水平下的VLM调用比较。右图：VLM调用缩减百分比，展示跨场景的鲁棒性。',
    '累积消融研究：逐步添加采样、级联过滤、自适应采样和优先级调度后的VLM调用次数。绿色注释显示相对于上一配置的增量缩减。',
]


def set_run_font(run, east_asia='宋体', ascii_font='Times New Roman', size=None, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run._element.rPr.rFonts.set(qn('w:ascii'), ascii_font)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), ascii_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def rewrite_paragraph(paragraph, text, east_asia='宋体', size=12, bold=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, size=size, bold=bold)
    return run


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_after(paragraph, text='', style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def run_pandoc(reference: Path, output: Path):
    pandoc = shutil.which('pandoc')
    if pandoc is None:
        raise RuntimeError('pandoc not found on PATH')

    command = [
        pandoc,
        str(PAPER_DIR / 'main.tex'),
        '--from=latex',
        '--to=docx',
        '--citeproc',
        f'--csl={PAPER_DIR / "numeric.csl"}',
        f'--bibliography={PAPER_DIR / "references.bib"}',
        '--metadata=reference-section-title:参考文献',
        f'--reference-doc={reference}',
        f'--resource-path={PAPER_DIR}',
        '-o',
        str(output),
    ]
    subprocess.run(command, check=True, cwd=PAPER_DIR.parent)


def copy_page_setup(doc, reference_doc):
    ref_section = reference_doc.sections[0]
    for section in doc.sections:
        section.page_width = ref_section.page_width
        section.page_height = ref_section.page_height
        section.top_margin = ref_section.top_margin
        section.bottom_margin = ref_section.bottom_margin
        section.left_margin = ref_section.left_margin
        section.right_margin = ref_section.right_margin


def style_title_block(doc):
    title = doc.paragraphs[0]
    title.style = doc.styles['Normal']
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(15.6)
    title.paragraph_format.space_after = Pt(15.6)
    rewrite_paragraph(title, TITLE, east_asia='黑体', size=18, bold=True)

    author = doc.paragraphs[1]
    author.style = doc.styles['Normal']
    author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(15.6)
    author.paragraph_format.space_after = Pt(15.6)
    rewrite_paragraph(author, AUTHOR, east_asia='楷体', size=12)


def style_abstracts(doc):
    paragraphs = doc.paragraphs
    if len(paragraphs) > 2 and paragraphs[2].text.strip() == 'Abstract':
        remove_paragraph(paragraphs[2])

    paragraphs = doc.paragraphs
    abstract = paragraphs[2]
    text = abstract.text.strip()
    if not text.startswith('摘要：'):
        text = '摘要：' + text
    abstract.style = doc.styles['Normal']
    abstract.paragraph_format.first_line_indent = Cm(0.74)
    abstract.paragraph_format.line_spacing = 1.5
    abstract.paragraph_format.space_before = Pt(0)
    abstract.paragraph_format.space_after = Pt(0)
    rewrite_paragraph(abstract, text, east_asia='楷体', size=12)

    keywords = paragraphs[3]
    keywords_text = keywords.text.strip().replace('关键词： ', '关键词：')
    keywords.style = doc.styles['Normal']
    keywords.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    keywords.paragraph_format.first_line_indent = None
    keywords.paragraph_format.line_spacing = 1.5
    rewrite_paragraph(keywords, keywords_text, east_asia='楷体', size=12)

    if not any(p.text.strip().startswith('Abstract:') for p in doc.paragraphs[:10]):
        en_abs = insert_paragraph_after(keywords, EN_ABSTRACT, doc.styles['Normal'])
        en_abs.paragraph_format.first_line_indent = Cm(0.74)
        en_abs.paragraph_format.line_spacing = 1.5
        for run in en_abs.runs:
            set_run_font(run, east_asia='楷体', size=12)
        en_key = insert_paragraph_after(en_abs, EN_KEYWORDS, doc.styles['Normal'])
        en_key.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        en_key.paragraph_format.first_line_indent = None
        en_key.paragraph_format.line_spacing = 1.5
        for run in en_key.runs:
            set_run_font(run, east_asia='楷体', size=12)


def fix_pandoc_text_loss(doc):
    replacements = {
        '自适应采样策略：': (
            '自适应采样策略：基于运动感知的可变捕获速率（活跃500ms / 空闲5000ms），'
            '在静态场景中减少约80%的帧数量。该策略通过像素级灰度差异阈值'
            '（θ=0.05）实现场景活动的自动检测。'
        ),
        '预算化本地VLM监控问题。': (
            '预算化本地VLM监控问题。给定摄像头帧流、轻量级观测函数m_t（运动）'
            '和o_t（目标标签），系统需要在有限VLM预算下选择帧子集发送给VLM。'
            '每次VLM调用成本约为t_v≈800ms，单GPU可持续吞吐近似为'
            '1/t_v≈1.25次/秒。调度目标是在不超过预算的前提下优先覆盖高优先级目标，'
            '并用固定回退间隔F=10s限制无目标场景的最长未分析时间。本文评估的是该'
            '预算化调度策略在oracle仿真条件下的调用量上界，而不是最优调度算法或'
            '真实风险检测准确率。'
        ),
        '阶段1：自适应帧捕获。': (
            '阶段1：自适应帧捕获。我们以可配置的间隔从视频流中捕获帧，并用运动'
            '检测调整捕获速率。系统维护活跃和空闲两个间隔，分别为500ms和5000ms。'
            '运动检测在小画布（160×120像素）上运行以最小化开销，计算像素级灰度差异。'
        ),
        '其中是像素的灰度值': (
            '其中I_t(i)是像素i在时刻t的灰度值，δ是像素阈值，运动比例超过5%时'
            '触发活跃模式。连续6帧无变化后，系统切换到空闲模式。'
        ),
        'detections , lastVlmTime, inFlight': (
            '算法1 基于优先级的VLM调度：输入detections、lastVlmTime和inFlight。'
            '若检测到person，则中止过时请求并立即调度VLM(frame)，同时更新lastVlmTime；'
            '否则，若检测到目标且无进行中请求，则调度VLM(frame)；否则，若'
            'now-lastVlmTime≥fallbackInterval且无进行中请求，则执行回退调度。'
        ),
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for prefix, replacement in replacements.items():
            if text.startswith(prefix):
                rewrite_paragraph(paragraph, replacement, size=12)
                break
        if '算法\xa0[alg:priority]' in paragraph.text:
            paragraph.text = paragraph.text.replace('算法\xa0[alg:priority]', '算法1')
            for run in paragraph.runs:
                set_run_font(run, size=12)


def normalize_caption(text):
    return text.replace('\\xa0', ' ').strip()


def apply_paragraph_styles(doc):
    table_numbers = {caption: i + 1 for i, caption in enumerate(TABLE_CAPTIONS)}
    figure_numbers = {caption: i + 1 for i, caption in enumerate(FIGURE_CAPTIONS)}
    in_references = False

    for idx, paragraph in enumerate(doc.paragraphs):
        text = normalize_caption(paragraph.text)
        if not text:
            continue
        if idx < 4 or text.startswith('Abstract:') or text.startswith('Keywords:'):
            continue

        if text in SECTION_TITLES or paragraph.style.name == 'Heading 1':
            paragraph.style = doc.styles['11']
            paragraph.paragraph_format.alignment = None
            in_references = text == '参考文献'
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=True)
            continue

        if text in SUBSECTION_TITLES:
            paragraph.style = doc.styles['22']
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=True)
            continue

        if text in table_numbers:
            paragraph.style = doc.styles['图标题']
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rewrite_paragraph(paragraph, f'表{table_numbers[text]} {text}', size=10.5)
            continue

        if text in figure_numbers:
            paragraph.style = doc.styles['图标题']
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rewrite_paragraph(paragraph, f'图{figure_numbers[text]} {text}', size=10.5)
            continue

        paragraph.style = doc.styles['4']
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = None if in_references else Cm(0.74)
        font_size = 10.5 if in_references else 12
        for run in paragraph.runs:
            set_run_font(run, size=font_size)


def remove_empty_tables(doc):
    for table in list(doc.tables):
        if len(table.rows) == 0:
            element = table._element
            element.getparent().remove(element)


def style_tables(doc):
    for table in doc.tables:
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        set_run_font(run, size=10.5)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--reference', type=Path, default=DEFAULT_REFERENCE)
    parser.add_argument('--output', type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    if not args.reference.exists():
        parser.error(
            'reference DOCX not found; pass --reference or set PAPER_REFERENCE_DOCX'
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        raw_output = Path(tmpdir) / 'main_123_format_raw.docx'
        run_pandoc(args.reference, raw_output)

        doc = Document(raw_output)
        reference_doc = Document(args.reference)
        copy_page_setup(doc, reference_doc)
        remove_empty_tables(doc)
        style_title_block(doc)
        style_abstracts(doc)
        fix_pandoc_text_loss(doc)
        apply_paragraph_styles(doc)
        style_tables(doc)
        doc.save(args.output)
    print(f'Word document saved to: {args.output}')


if __name__ == '__main__':
    main()
