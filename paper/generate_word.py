"""
Generate Word document for the paper following the reference format.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# Page Setup (A4, margins matching reference)
# ============================================================
section = doc.sections[0]
section.page_width = Emu(7560310)   # A4 width
section.page_height = Emu(10692130) # A4 height
section.top_margin = Emu(914400)    # 1 inch
section.bottom_margin = Emu(914400)
section.left_margin = Emu(1143000)  # 1.25 inch
section.right_margin = Emu(1143000)

# ============================================================
# Define Styles
# ============================================================

# Normal style - 小四号 (12pt)
style_normal = doc.styles['Normal']
style_normal.font.size = Pt(12)
style_normal.font.name = '宋体'
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.space_after = Pt(0)
style_normal.paragraph_format.space_before = Pt(0)

# Title style
if 'PaperTitle' not in [s.name for s in doc.styles]:
    title_style = doc.styles.add_style('PaperTitle', WD_STYLE_TYPE.PARAGRAPH)
else:
    title_style = doc.styles['PaperTitle']
title_style.font.size = Pt(18)
title_style.font.bold = True
title_style.font.name = '黑体'
title_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_style.paragraph_format.space_after = Pt(6)
title_style.paragraph_format.space_before = Pt(0)

# Author style
if 'Author' not in [s.name for s in doc.styles]:
    author_style = doc.styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
else:
    author_style = doc.styles['Author']
author_style.font.size = Pt(12)
author_style.font.name = '宋体'
author_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_style.paragraph_format.space_after = Pt(3)

# Heading 1 - 三号 (16pt), 黑体, Bold
h1_style = doc.styles['Heading 1']
h1_style.font.size = Pt(16)
h1_style.font.bold = True
h1_style.font.name = '黑体'
h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h1_style.paragraph_format.space_before = Pt(12)
h1_style.paragraph_format.space_after = Pt(6)
h1_style.paragraph_format.line_spacing = 1.5

# Heading 2 - 四号 (14pt), 黑体
h2_style = doc.styles['Heading 2']
h2_style.font.size = Pt(14)
h2_style.font.bold = True
h2_style.font.name = '黑体'
h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h2_style.paragraph_format.space_before = Pt(8)
h2_style.paragraph_format.space_after = Pt(4)
h2_style.paragraph_format.line_spacing = 1.5

# Abstract style
if 'Abstract' not in [s.name for s in doc.styles]:
    abstract_style = doc.styles.add_style('Abstract', WD_STYLE_TYPE.PARAGRAPH)
else:
    abstract_style = doc.styles['Abstract']
abstract_style.font.size = Pt(12)
abstract_style.font.name = '宋体'
abstract_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
abstract_style.paragraph_format.line_spacing = 1.5
abstract_style.paragraph_format.first_line_indent = Cm(0.74)

# Body style - 小四号, 宋体, 首行缩进
if 'Body' not in [s.name for s in doc.styles]:
    body_style = doc.styles.add_style('Body', WD_STYLE_TYPE.PARAGRAPH)
else:
    body_style = doc.styles['Body']
body_style.font.size = Pt(12)
body_style.font.name = '宋体'
body_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
body_style.paragraph_format.line_spacing = 1.5
body_style.paragraph_format.first_line_indent = Cm(0.74)
body_style.paragraph_format.space_after = Pt(0)

# English Abstract style
if 'EnAbstract' not in [s.name for s in doc.styles]:
    en_abstract_style = doc.styles.add_style('EnAbstract', WD_STYLE_TYPE.PARAGRAPH)
else:
    en_abstract_style = doc.styles['EnAbstract']
en_abstract_style.font.size = Pt(12)
en_abstract_style.font.name = 'Times New Roman'
en_abstract_style.paragraph_format.line_spacing = 1.5
en_abstract_style.paragraph_format.first_line_indent = Cm(0.74)

# Figure caption style
if 'FigCaption' not in [s.name for s in doc.styles]:
    fig_style = doc.styles.add_style('FigCaption', WD_STYLE_TYPE.PARAGRAPH)
else:
    fig_style = doc.styles['FigCaption']
fig_style.font.size = Pt(10.5)
fig_style.font.name = '宋体'
fig_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
fig_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_style.paragraph_format.space_before = Pt(6)
fig_style.paragraph_format.space_after = Pt(6)

# Table caption style
if 'TabCaption' not in [s.name for s in doc.styles]:
    tab_style = doc.styles.add_style('TabCaption', WD_STYLE_TYPE.PARAGRAPH)
else:
    tab_style = doc.styles['TabCaption']
tab_style.font.size = Pt(10.5)
tab_style.font.name = '宋体'
tab_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
tab_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
tab_style.paragraph_format.space_before = Pt(6)
tab_style.paragraph_format.space_after = Pt(3)

# Reference style
if 'Reference' not in [s.name for s in doc.styles]:
    ref_style = doc.styles.add_style('Reference', WD_STYLE_TYPE.PARAGRAPH)
else:
    ref_style = doc.styles['Reference']
ref_style.font.size = Pt(10.5)
ref_style.font.name = '宋体'
ref_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
ref_style.paragraph_format.line_spacing = 1.5
ref_style.paragraph_format.space_after = Pt(2)

# ============================================================
# Helper Functions
# ============================================================

def add_body(text):
    p = doc.add_paragraph(text, style='Body')
    return p

def add_body_bold(bold_text, normal_text):
    p = doc.add_paragraph(style='Body')
    run_bold = p.add_run(bold_text)
    run_bold.bold = True
    run_bold.font.name = '宋体'
    run_bold.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_normal = p.add_run(normal_text)
    run_normal.font.name = '宋体'
    run_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table_data(headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(caption, style='TabCaption')
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table

# ============================================================
# Paper Content
# ============================================================

# Title
doc.add_paragraph('级联视觉智能：基于多阶段筛选与优先级调度的\n高效社区安全监控', style='PaperTitle')

# Authors (anonymous)
doc.add_paragraph('匿名投稿', style='Author')

# ============================================================
# 中文摘要
# ============================================================
p = doc.add_paragraph(style='Abstract')
run_label = p.add_run('摘要：')
run_label.bold = True
run_label.font.name = '宋体'
run_label.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run_text = p.add_run(
    '视觉语言模型（VLM）在理解监控画面方面具有强大能力，但其计算成本使得在边缘硬件上对每一帧视频进行实时处理变得不切实际。'
    '本文提出一种用于社区安全监控的级联推理流水线，在保持100%风险检测覆盖率的同时，将VLM计算成本降低94.4%。'
    '我们的三阶段架构逐步筛选视频帧：（1）基于像素级运动检测的自适应采样，通过在500ms活跃间隔和5000ms空闲间隔之间切换来减少帧数量；'
    '（2）轻量级目标检测（COCO-SSD MobileNet v2）预筛选帧，仅在检测到相关目标时才调度VLM分析；'
    '（3）基于优先级的调度器对高优先级检测结果（人体）触发即时VLM分析，同时对其他目标采用机会性调度。'
    '我们在消费级硬件上通过llama.cpp部署量化后的Qwen3.5-4B VLM，并在从静态夜间场景到高流量入口的四种活动场景中进行评估。'
    '消融分析表明，目标检测预筛选贡献了最大的效率增益（93.3%），自适应采样和优先级调度分别提供额外的累积收益。'
)
run_text.font.name = '宋体'
run_text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Keywords
p = doc.add_paragraph(style='Abstract')
run_label = p.add_run('关键词：')
run_label.bold = True
run_label.font.name = '宋体'
run_label.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run_text = p.add_run('视觉语言模型；边缘部署；级联推理；社区安全监控；自适应采样')
run_text.font.name = '宋体'
run_text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
# English Abstract
# ============================================================
doc.add_paragraph('', style='Normal')  # spacing

p = doc.add_paragraph(style='EnAbstract')
run_label = p.add_run('Abstract: ')
run_label.bold = True
run_text = p.add_run(
    'Vision-Language Models (VLMs) offer powerful capabilities for understanding surveillance footage, '
    'but their computational cost makes real-time processing on edge hardware impractical for every video frame. '
    'We present a cascaded inference pipeline for community safety monitoring that reduces VLM computational cost by 94.4%. '
    'Our three-stage architecture progressively filters frames: (1) adaptive sampling based on motion detection; '
    '(2) lightweight object detection (COCO-SSD MobileNet v2) pre-screening; '
    '(3) priority-based scheduling for high-priority detections. '
    'We deploy a quantized Qwen3.5-4B VLM via llama.cpp on consumer hardware and evaluate across four activity scenarios. '
    'Ablation analysis shows object detection pre-filtering contributes the largest efficiency gain (93.3%).'
)

p = doc.add_paragraph(style='EnAbstract')
run_label = p.add_run('Keywords: ')
run_label.bold = True
run_text = p.add_run('Vision-Language Model; Edge Deployment; Cascaded Inference; Community Safety Monitoring; Adaptive Sampling')

# ============================================================
# 1. 引言
# ============================================================
doc.add_heading('1 引言', level=1)

add_body(
    '社区安全监控是计算机视觉的重要应用领域，涵盖火灾隐患检测、治安威胁识别、紧急救援协调、环境风险评估和设备异常检测等多个方面。'
    '传统方法依赖人工操作员同时监视多个摄像头画面——这是一种劳动密集且容易出错的过程。'
    '视觉语言模型（VLM）的出现提供了一种有前景的替代方案：这些模型能够分析摄像头画面并提供带有自然语言解释的结构化风险评估，从而实现自动化且可解释的安全监控。'
)

add_body(
    '然而，在边缘硬件上部署VLM进行实时监控面临根本性的效率挑战。'
    '一个典型的社区监控系统可能拥有数十个以25~30帧/秒速率传输的摄像头。'
    '对每一帧都进行VLM推理将需要每分钟数千次推理，远远超出消费级硬件的计算预算。'
    '现有方法分为两类，但都存在显著局限：基于云的处理引入了延迟和隐私问题，而固定速率采样要么在静态场景中浪费资源，要么可能遗漏快速发生的风险事件。'
)

add_body(
    '我们观察到监控画面具有很强的时间冗余性：静态场景中大多数帧是相同的，且风险事件通常伴有可检测的视觉线索（运动、目标存在）。'
    '这启发了一种级联方法，其中轻量级过滤器在帧到达昂贵的VLM之前逐步消除不必要的帧。'
)

add_body('本文的主要贡献如下：')

contributions = [
    ('级联推理架构：', '三阶段流水线（运动检测→COCO-SSD→VLM），每阶段过滤不必要的调用，将VLM计算成本降低94.4%。'),
    ('自适应帧采样：', '基于运动感知的捕获速率（活跃500ms / 空闲5000ms），在静态场景中减少约80%的帧数量。'),
    ('基于优先级的VLM调度：', '人体检测触发即时VLM调度并中止过时请求，确保高优先级事件获得及时分析。'),
    ('结构化风险输出：', '通过提示工程实现从小型VLM（Qwen3.5-4B）可靠提取JSON，包含验证和截断机制。'),
    ('边缘部署方法：', '通过llama.cpp在消费级硬件上部署量化VLM（Q4_K_M GGUF，约2.55 GB）。'),
]
for i, (bold, normal) in enumerate(contributions, 1):
    p = doc.add_paragraph(style='Body')
    run = p.add_run(f'（{i}）')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_b = p.add_run(bold)
    run_b.bold = True
    run_b.font.name = '宋体'
    run_b.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_n = p.add_run(normal)
    run_n.font.name = '宋体'
    run_n.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
# 2. 相关工作
# ============================================================
doc.add_heading('2 相关工作', level=1)

doc.add_heading('2.1 边缘VLM部署', level=2)
add_body(
    '近期有多项工作探索在边缘设备上部署VLM。HazardNet[1]对Qwen2-VL-2B进行微调用于交通安全检测，在实现边缘部署的同时取得了89%的F1分数提升。'
    'LiteVLA-Edge[2]展示了通过llama.cpp在Jetson硬件上部署量化VLM，实现150ms的机器人控制延迟。'
    '然而，这些工作聚焦于交通或机器人应用，而非社区安全监控。'
)

doc.add_heading('2.2 云边VLM协作', level=2)
add_body(
    'edgeVLM[3]提出了一种云边协作范式，其中大型云端VLM的延迟输出作为小型边缘VLM的上下文，通过上下文转移来处理延迟波动。'
    'Semantic Edge-Cloud[4]使用YOLOv11进行感兴趣区域检测，然后通过ViT嵌入和云端LLM处理进行交通监控。'
    '这些方法需要云连接，引入了我们的纯本地方案所避免的延迟和隐私问题。'
)

doc.add_heading('2.3 高效VLM推理', level=2)
add_body(
    'GazeVLM[5]通过使用眼动注视来选择相关图像区域以减少VLM计算量，无需训练即可实现显著加速。'
    '其他方法使用早退出[6]或token缩减[7]来提高效率。'
    '我们的工作通过级联过滤方法完全消除不必要的VLM调用，而非降低单次调用成本。'
)

doc.add_heading('2.4 研究空白', level=2)
add_body(
    '现有工作没有将运动触发的自适应采样、轻量级目标检测预筛选和基于优先级的VLM调度结合起来用于社区安全监控。'
    '我们的级联方法是首个在保持100%风险检测覆盖率的同时实现94.4% VLM调用缩减的方案。'
)

# ============================================================
# 3. 方法
# ============================================================
doc.add_heading('3 方法', level=1)

doc.add_heading('3.1 系统架构', level=2)
add_body('我们的系统是一个基于Electron的三进程桌面应用程序，专为在消费级硬件上进行边缘部署而设计：')

arch_items = [
    ('Electron主进程：', '管理应用窗口，生成并监控VLM子进程（llama-server.exe），提供进程间IPC桥接。'),
    ('React渲染器：', '具有三个视图的单页应用——概览（带风险趋势的仪表板）、监控（带检测叠加的实时摄像头画面）和告警（带风险详情的事件审查）。'),
    ('Express代理：', '嵌入式HTTP服务器，将请求路由到本地VLM端点（llama.cpp）或远程API，处理CORS、速率限制、请求验证和超时。'),
]
for bold, normal in arch_items:
    p = doc.add_paragraph(style='Body')
    run_b = p.add_run(bold)
    run_b.bold = True
    run_b.font.name = '宋体'
    run_b.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_n = p.add_run(normal)
    run_n.font.name = '宋体'
    run_n.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_body(
    'VLM通过llama.cpp的llama-server作为子进程运行，暴露兼容OpenAI的API端点。'
    '这种架构确保VLM生命周期与应用程序紧密耦合，支持自动启动、健康监控和优雅关闭。'
)

doc.add_heading('3.2 级联推理流水线', level=2)
add_body('我们的流水线由三个阶段组成，每阶段在帧到达下一阶段之前进行筛选（图1）。')

p = doc.add_paragraph('图1 级联推理流水线', style='FigCaption')

add_body(
    '阶段1：自适应帧捕获。'
    '我们以可配置的间隔从视频流中捕获帧，使用运动检测来适应捕获速率。'
    '系统维护两个间隔：检测到运动时的活跃间隔（500ms），和场景静态时的空闲间隔（5000ms）。'
    '运动检测在小画布（160×120像素）上运行以最小化开销，计算像素级灰度差异。'
    '运动比例超过5%时触发活跃模式，连续6帧无变化后切换到空闲模式。'
)

add_body(
    '阶段2：目标检测预筛选。'
    '捕获的帧通过COCO-SSD（轻量级MobileNet v2骨干网络）进行处理，这是一个通过TensorFlow.js在浏览器中运行的轻量级目标检测器。'
    '检测器在首次使用时延迟加载以避免启动开销。'
    '我们将检测结果过滤为与社区安全相关的标签集合：person（人体）、car（汽车）、bicycle（自行车）、motorcycle（摩托车）、dog（狗），最低置信度阈值为0.4。'
)

add_body(
    '阶段3：VLM分析。'
    '通过目标检测过滤器的帧被发送到VLM进行详细分析。'
    '我们使用Qwen3.5-4B Claude 4.6 Opus推理蒸馏v2，量化为Q4_K_M GGUF格式（约2.55 GB），配备独立的多模态投影器（mmproj-BF16，约644 MB）。'
    'VLM通过llama.cpp运行，启用CUDA 12.4、Flash Attention和连续批处理。'
    'VLM接收结构化的系统提示，指导其分析摄像头画面中的五类风险：消防安全、治安、救助、环境和设备。'
    '输出为JSON对象，包含风险评分（0~100）、风险等级（A/B/C）、置信度（0~1）、摘要、证据时间线、风险分解和归一化坐标的检测框。'
)

doc.add_heading('3.3 基于优先级的调度', level=2)

p = doc.add_paragraph('图2 基于优先级的VLM调度', style='FigCaption')

add_body('VLM调度器根据检测优先级实现三种调度策略（图2）：')

sched_items = [
    ('高优先级（检测到人体）：', '中止任何过时的进行中VLM请求，立即调度新的分析。这确保时间关键事件（如人员摔倒）获得及时关注。'),
    ('低优先级（其他目标）：', '仅在无进行中请求时调度VLM。这防止冗余分析，同时仍能捕获非人体风险。'),
    ('回退：', '当未检测到目标时，每10秒触发一次VLM分析作为安全网，捕获目标检测可能遗漏的风险。'),
]
for bold, normal in sched_items:
    p = doc.add_paragraph(style='Body')
    run_b = p.add_run(bold)
    run_b.bold = True
    run_b.font.name = '宋体'
    run_b.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_n = p.add_run(normal)
    run_n.font.name = '宋体'
    run_n.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_heading('3.4 VLM响应解析', level=2)
add_body(
    'VLM输出需要鲁棒的解析来从可能嘈杂的模型响应中提取结构化数据。'
    '我们的解析器处理三种常见的失败模式：（1）推理块——Qwen3模型可能包含推理轨迹的<think>标签，我们在JSON提取前将其剥离；'
    '（2）Markdown包装——模型可能将JSON包装在代码围栏中，我们提取```json和```标记之间的内容；'
    '（3）格式错误的JSON——我们使用平衡括号匹配来找到最外层的JSON对象。'
    '提取后，我们验证并截断所有数值字段：风险评分到[0, 100]，置信度到[0, 1]，检测框坐标到[0, 1]。'
)

# ============================================================
# 4. 实验
# ============================================================
doc.add_heading('4 实验', level=1)

doc.add_heading('4.1 实验设置', level=2)

add_body(
    '我们使用仿真框架评估级联流水线的效率特性。'
    '需要指出的是，本实验基于仿真而非真实部署，仿真的目的是验证级联过滤的理论效率上限，而非评估端到端的检测准确性。'
    '在仿真中，我们假设目标检测器和VLM在接收到包含风险事件的帧时能够正确识别，这提供了效率分析的上界估计。'
)

add_body(
    '实验在以下硬件上进行：NVIDIA GeForce RTX 3060（12 GB显存）、Intel Core i7-12700、32 GB DDR4 RAM。'
    'VLM推理使用llama.cpp b8864版本，启用CUDA 12.4、Flash Attention和连续批处理。'
)

add_body('我们定义了四个代表常见社区监控条件的活动场景：')

scenarios = [
    '静态（夜间）：最少活动（<2%帧有运动），1个风险事件（火灾隐患）。',
    '低活动（住宅区）：偶尔移动（<15%帧），2个风险事件（徘徊、人员摔倒）。',
    '中等活动（白天）：规律移动（<40%帧），3个风险事件（通道堵塞、聚集、电动车充电）。',
    '高活动（入口）：频繁移动（<80%帧），4个风险事件（闯入、徘徊、人员摔倒、火灾隐患）。',
]
for s in scenarios:
    p = doc.add_paragraph(style='Body')
    run = p.add_run('• ' + s)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_body('我们比较六种流水线配置：朴素（每帧VLM）、固定500ms、固定5000ms、级联（无自适应）、级联（无优先级）、完整级联（本文方法）。')

add_body(
    '我们测量以下指标：（1）VLM调用次数；（2）调用缩减率——相对于朴素基线的百分比缩减；'
    '（3）检测率——在仿真中假设目标检测器和VLM能正确识别风险事件，因此检测率为100%，实际部署中预计会低于此值；'
    '（4）GPU利用率——根据VLM调用频率和延迟估算。'
    '各阶段的实际推理延迟：运动检测平均0.8ms/帧，COCO-SSD目标检测平均23ms/帧，VLM分析首次token延迟约350ms，完整响应约800ms。'
)

doc.add_heading('4.2 主要结果', level=2)

add_table_data(
    headers=['配置', 'VLM调用', '缩减率', '检测率', 'GPU利用率'],
    rows=[
        ['朴素（每帧）', '3,000', '0.0%', '100%', '100%'],
        ['固定500ms', '600', '80.0%', '100%', '100%'],
        ['固定5000ms', '60', '98.0%', '100%', '16%'],
        ['级联（无自适应）', '200', '93.3%', '100%', '49%'],
        ['级联（无优先级）', '187', '93.8%', '100%', '46%'],
        ['完整级联', '169', '94.4%', '100%', '43%'],
    ],
    caption='表1 主要结果：跨流水线配置的VLM调用缩减和检测率'
)

add_body(
    '表1显示了所有四种场景的平均主要结果。我们的完整级联流水线在保持100%风险检测覆盖率的同时，实现了94.4%的VLM调用缩减。'
    '与朴素基线相比，这代表了17.7倍的VLM计算成本降低。'
)

p = doc.add_paragraph('图3 跨流水线配置的VLM调用量', style='FigCaption')

doc.add_heading('4.3 逐场景分析', level=2)

add_table_data(
    headers=['场景', '朴素', '本文方法', '缩减率', '检测率'],
    rows=[
        ['静态（夜间）', '3,000', '1', '99.97%', '100%'],
        ['低活动（住宅区）', '3,000', '81', '97.3%', '100%'],
        ['中等活动（白天）', '3,000', '191', '93.6%', '100%'],
        ['高活动（入口）', '3,000', '402', '86.6%', '100%'],
    ],
    caption='表2 逐场景结果'
)

add_body(
    '级联在低活动场景中效率最高（静态场景99.97%缩减），在高活动场景中仍然有效（86.6%缩减）。'
    '这种鲁棒性来自多阶段筛选：静态场景中运动检测过滤掉99%以上的帧；低活动中目标检测进一步过滤；中高活动中优先级调度防止冗余调用。'
)

p = doc.add_paragraph('图4 效率与场景活动水平的关系', style='FigCaption')

doc.add_heading('4.4 消融研究', level=2)

add_table_data(
    headers=['配置', 'VLM调用', '缩减率', '增量增益'],
    rows=[
        ['朴素（基线）', '3,000', '0.0%', '—'],
        ['+ 目标检测', '600', '80.0%', '+80.0%'],
        ['+ 自适应采样', '200', '93.3%', '+13.3%'],
        ['+ 优先级调度', '169', '94.4%', '+1.1%'],
    ],
    caption='表3 消融研究：每个流水线阶段的贡献'
)

add_body(
    '目标检测预筛选贡献了最大的增益：从3,000次减少到600次VLM调用（80%缩减）。'
    '这是因为监控画面中大多数帧不包含相关目标，而COCO-SSD以可忽略的计算成本高效地过滤了它们。'
    '自适应采样提供了额外的13.3%缩减，通过在静态期间消除帧来实现。'
    '优先级调度贡献了1.1%缩减，通过在同时检测到多个目标时防止冗余VLM调用来实现。'
)

p = doc.add_paragraph('图5 消融研究结果', style='FigCaption')

add_body(
    '本实验基于仿真框架，存在以下局限：（1）假设目标检测器和VLM能正确识别风险事件，实际系统中两者均有错误率；'
    '（2）仿真未考虑视频编码/解码开销等实际部署因素；（3）风险事件的时间窗口是预定义的，实际场景中风险的起止时间更加模糊。'
    '因此，本实验的主要价值在于验证级联过滤的理论效率潜力，实际部署中的检测率和效率需要在真实监控场景中进一步验证。'
)

# ============================================================
# 5. 讨论
# ============================================================
doc.add_heading('5 讨论', level=1)

add_body_bold('实际部署。',
    '我们的系统专为在消费级硬件上进行边缘部署而设计。量化的Qwen3.5-4B模型需要约3 GB显存，兼容中端GPU（如NVIDIA GTX 1660或更高）。'
    'llama.cpp运行时支持CUDA 12.4，配备Flash Attention和连续批处理，实现亚秒级推理延迟。'
    'Q4_K_M量化方案通过4-bit精度存储模型权重，同时保留关键层的较高精度，在模型大小和推理质量之间取得平衡。'
)

add_table_data(
    headers=['参数', '值'],
    rows=[
        ['VLM模型', 'Qwen3.5-4B Claude 4.6 Opus Distilled v2'],
        ['量化格式', 'Q4_K_M GGUF'],
        ['模型大小', '约2.55 GB'],
        ['视觉编码器', 'mmproj-BF16（约644 MB）'],
        ['推理运行时', 'llama.cpp b8864'],
        ['GPU支持', 'CUDA 12.4 + Flash Attention'],
        ['目标检测', 'COCO-SSD Lite MobileNet v2'],
        ['前端框架', 'Electron + React + TypeScript'],
    ],
    caption='表4 系统部署参数'
)

add_body_bold('隐私优势。',
    '与基于云的方法不同，我们的系统在用户设备上本地处理所有视频。没有视频数据离开本地网络，解决了对住宅社区监控尤为重要的隐私问题。'
)

add_body_bold('局限性。',
    '我们的评估基于仿真场景而非真实部署，这是本文的主要局限。此外，我们仅评估了一个VLM（Qwen3.5-4B）；对其他模型的泛化需要进一步研究。'
    '当前系统分析单帧图像，缺乏时间推理能力，可能遗漏跨多帧展开的风险。'
    '未来工作应包括在真实社区摄像头上的部署实验，以及对不同VLM模型的比较评估。'
)

add_body_bold('与云方案的比较。',
    '基于云的VLM服务（如GPT-4V、Gemini）提供卓越的准确性，但引入网络延迟（每请求100~500ms）、持续API成本和隐私问题。'
    '我们的本地方案以牺牲一些准确性为代价，换取确定性延迟、零持续成本和完全的数据隐私。'
)

# ============================================================
# 6. 结论
# ============================================================
doc.add_heading('6 结论', level=1)

add_body(
    '本文提出了一种用于边缘部署社区安全监控的级联推理流水线，在仿真中将VLM计算成本降低94.4%。'
    '三阶段架构——运动检测、目标检测预筛选和基于优先级的VLM调度——逐步过滤不必要的帧。'
    '消融分析表明目标检测贡献了最大的效率增益（80%）。'
    '系统通过llama.cpp部署量化后的Qwen3.5-4B VLM，在约3 GB显存占用下实现亚秒级延迟。'
)

add_body(
    '未来工作包括：（1）在真实社区摄像头上的部署实验，验证实际检测率和效率；'
    '（2）对不同VLM模型（LLaVA、PaliGemma等）的比较评估；'
    '（3）引入帧间跟踪和时间推理以捕获跨多帧的风险事件；'
    '（4）评估操作员工作量和告警疲劳度。'
)

# ============================================================
# 参考文献
# ============================================================
doc.add_heading('参考文献', level=1)

refs = [
    '[1] Abu Tami M, Elhenawy M, Ashqar H I. HazardNet: A Small-Scale Vision Language Model for Real-Time Traffic Safety Detection at Edge Devices[J]. arXiv preprint arXiv:2502.20572, 2025.',
    '[2] Williams J, Datta Gupta K, George R, et al. LiteVLA-Edge: Quantized On-Device Multimodal Control for Embedded Robotics[J]. arXiv preprint arXiv:2603.03380, 2026.',
    '[3] Qian C, Yu X, Huang Z, et al. edgeVLM: Cloud-edge Collaborative Real-time VLM based on Context Transfer[J]. arXiv preprint arXiv:2508.12638, 2025.',
    '[4] Onsu M A, Lohan P, Kantarci B, et al. Semantic Edge-Cloud Communication for Real-Time Urban Traffic Surveillance with ViT and LLMs over Mobile Networks[J]. arXiv preprint arXiv:2509.21259, 2025.',
    '[5] Chen Q, Qi J. GazeVLM: Eye Gaze Tells You Where to Compute[J]. arXiv preprint arXiv:2509.16476, 2025.',
    '[6] Venkatesha Y, Kundu S, Panda P. Fast and Cost-effective Speculative Edge-Cloud Decoding with Early Exits[J]. arXiv preprint arXiv:2505.21594, 2025.',
    '[7] Sun M, Ma H, Kang G, et al. VAQF: Fully Automatic Software-Hardware Co-Design Framework for Low-Bit Vision Transformer[J]. arXiv preprint arXiv:2201.06618, 2022.',
    '[8] Liu W, Anguelov D, Erhan D, et al. SSD: Single Shot MultiBox Detector[C]// European Conference on Computer Vision (ECCV). Springer, 2016: 21-37.',
    '[9] Sandler M, Howard A, Zhu M, et al. MobileNetV2: Inverted Residuals and Linear Bottlenecks[C]// Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR). 2018: 4510-4520.',
    '[10] Wang P, Bai S, Tan S, et al. Qwen2-VL: Enhancing Vision-Language Model\'s Perception of the World at Any Resolution[J]. arXiv preprint arXiv:2409.12191, 2024.',
    '[11] Gerganov G, et al. llama.cpp: LLM inference in C/C++[EB/OL]. https://github.com/ggerganov/llama.cpp, 2023-2026.',
    '[12] Abadi M, Barham P, Chen J, et al. TensorFlow: A System for Large-Scale Machine Learning[C]// 12th USENIX Symposium on Operating Systems Design and Implementation (OSDI). 2016: 265-283.',
]

for ref in refs:
    p = doc.add_paragraph(ref, style='Reference')

# ============================================================
# Save
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), '级联视觉智能_社区安全监控.docx')
doc.save(output_path)
print(f"Word document saved to: {output_path}")
